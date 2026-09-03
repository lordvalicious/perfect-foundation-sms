"""
Generate Simple Perfect Foundation SMS Report.
Run: python generate_report.py
Output: Perfect_Foundation_SMS_Report_v2.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, "1a3c6e")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
    for i, data in enumerate(rows):
        for j, val in enumerate(data):
            table.rows[i + 1].cells[j].text = str(val)
    return table


def build_report():
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER ──
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_heading("Perfect Foundation SMS", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("Project Report", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for line in ["", "September 2026", "Version 2.0"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── TABLE OF CONTENTS ──
    doc.add_page_break()
    doc.add_heading("Contents", level=1)
    
    toc_items = [
        "1. Overview",
        "2. Technology Stack",
        "3. Project Structure",
        "4. Demo Data",
        "5. Deployment",
        "6. Test Results",
        "7. Security",
        "8. Known Issues",
        "9. Recommendations",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 1. Overview
    # ══════════════════════════════════════
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Perfect Foundation SMS is a school management system for multi-campus "
        "institutions. It handles students, staff, finances, academics, attendance, "
        "and communications with role-based access control."
    )

    doc.add_heading("Key Features", level=2)
    doc.add_paragraph(
        "• Multi-campus operations with data isolation\n"
        "• Student enrollment and tracking\n"
        "• Staff and teacher management\n"
        "• Invoice generation and payment tracking\n"
        "• Attendance and gradebook\n"
        "• Internal messaging and notifications\n"
        "• PDF report generation"
    )

    # ══════════════════════════════════════
    # 2. Technology Stack
    # ══════════════════════════════════════
    doc.add_heading("2. Technology Stack", level=1)
    add_table(doc, ["Layer", "Technology"], [
        ["Backend", "Django 6.1, Python 3.12"],
        ["API", "Django REST Framework 3.18"],
        ["Frontend", "React 19, Vite 5.4"],
        ["Database", "Neon PostgreSQL"],
        ["Hosting", "Vercel (serverless)"],
        ["Styling", "Tailwind CSS"],
    ])

    # ══════════════════════════════════════
    # 3. Project Structure
    # ══════════════════════════════════════
    doc.add_heading("3. Project Structure", level=1)
    doc.add_paragraph(
        "backend/\n"
        "  api/index.py          - Vercel serverless entry\n"
        "  apps/accounts/        - Auth, users, students, staff\n"
        "  apps/schools/         - Institutions, campuses\n"
        "  apps/finance/         - Invoices, payments\n"
        "  apps/academics/       - Courses, homework\n"
        "  apps/attendance/      - Attendance tracking\n"
        "  apps/communication/   - Messages, notifications\n"
        "  config/settings/      - Production, test, dev settings\n\n"
        "frontend/\n"
        "  src/pages/            - Application pages\n"
        "  src/services/         - API client\n"
        "  vercel.json           - Deployment config"
    )

    # ══════════════════════════════════════
    # 4. Demo Data
    # ══════════════════════════════════════
    doc.add_heading("4. Demo Data", level=1)
    doc.add_paragraph(
        "The system includes a demo data seeder with 5 parts. Run:\n"
        "  python manage.py seed_demo_data --all"
    )

    doc.add_heading("Seeded Data", level=2)
    add_table(doc, ["Entity", "Count"], [
        ["Students", "500"],
        ["Enrollments", "500"],
        ["Guardians", "350"],
        ["Teachers", "50"],
        ["Staff", "96"],
        ["Users", "1,016"],
        ["Campuses", "5"],
        ["Invoices", "500"],
        ["Payments", "443"],
        ["Courses", "25"],
    ])

    doc.add_heading("Demo Login", level=2)
    doc.add_paragraph(
        "Username: demo_superadmin\n"
        "Password: DemoPassword!2026"
    )

    # ══════════════════════════════════════
    # 5. Deployment
    # ══════════════════════════════════════
    doc.add_heading("5. Deployment", level=1)

    doc.add_heading("Live URLs", level=2)
    add_table(doc, ["Service", "URL"], [
        ["Frontend", "https://perfect-foundation-sms.vercel.app"],
        ["Backend API", "https://perfect-foundation-api.vercel.app"],
        ["Health Check", "https://perfect-foundation-sms.vercel.app/api/health/"],
    ])

    doc.add_heading("Backend Deployment", level=2)
    doc.add_paragraph(
        "• Entry point: api/index.py (ASGI)\n"
        "• Runtime: Python 3.12 serverless\n"
        "• Database: Neon PostgreSQL (neondb_owner role)\n"
        "• Build: Auto-installs requirements.txt\n"
        "• Migrations: Run via vercel.json buildCommand"
    )

    doc.add_heading("Frontend Deployment", level=2)
    doc.add_paragraph(
        "• Build: Vite static output\n"
        "• API proxy: /api/* → backend API\n"
        "• SPA fallback: All routes → index.html"
    )

    doc.add_heading("Environment Variables", level=2)
    add_table(doc, ["Variable", "Purpose"], [
        ["DATABASE_URL", "Neon PostgreSQL connection"],
        ["DJANGO_SECRET_KEY", "Production secret key"],
        ["DJANGO_ALLOWED_HOSTS", ".vercel.app, localhost"],
        ["DJANGO_SETTINGS_MODULE", "config.settings.production"],
    ])

    # ══════════════════════════════════════
    # 6. Test Results
    # ══════════════════════════════════════
    doc.add_heading("6. Test Results", level=1)
    
    add_table(doc, ["Metric", "Count"], [
        ["Total tests", "447"],
        ["Passing", "376"],
        ["Failing", "39"],
        ["Errors", "32"],
    ])

    doc.add_paragraph(
        "Main issues:\n"
        "• Rate limiting blocks auth tests (12 tests)\n"
        "• Pre-existing code bugs (25 tests)\n"
        "• Test file defects (5 tests)\n\n"
        "Note: None of these failures are caused by the demo seed system."
    )

    # ══════════════════════════════════════
    # 7. Security
    # ══════════════════════════════════════
    doc.add_heading("7. Security", level=1)
    doc.add_paragraph(
        "• JWT authentication with 14-day sessions\n"
        "• Role-based access control (Super Admin, Org Admin, Teacher, Student)\n"
        "• CSRF protection with cross-origin support\n"
        "• Rate limiting on login endpoints\n"
        "• HTTPS enforced in production\n"
        "• Passwords hashed with PBKDF2\n"
        "• Multi-tenant data isolation"
    )

    # ══════════════════════════════════════
    # 8. Known Issues
    # ══════════════════════════════════════
    doc.add_heading("8. Known Issues", level=1)
    doc.add_paragraph(
        "1. Duplicate SubjectOffering model in schools/models.py\n"
        "   Causes RuntimeWarning on startup.\n\n"
        "2. StudentTransfer.objects manager rebind (accounts/models.py:930)\n"
        "   Broken manager binding; workaround in place.\n\n"
        "3. Test rate limiting\n"
        "   Login throttles not fully disabled in test settings.\n\n"
        "4. Student360Serializer redundancy\n"
        "   Causes DRF assertion error."
    )

    # ══════════════════════════════════════
    # 9. Recommendations
    # ══════════════════════════════════════
    doc.add_heading("9. Recommendations", level=1)

    doc.add_heading("Immediate", level=2)
    doc.add_paragraph(
        "• Fix duplicate SubjectOffering model\n"
        "• Fix test rate limiting configuration\n"
        "• Add database indexes for performance"
    )

    doc.add_heading("Short-term", level=2)
    doc.add_paragraph(
        "• Add caching for dashboard stats\n"
        "• Implement API pagination\n"
        "• Add error monitoring (Sentry)"
    )

    doc.add_heading("Long-term", level=2)
    doc.add_paragraph(
        "• Mobile app (React Native)\n"
        "• Multi-language support\n"
        "• Advanced analytics\n"
        "• External LMS integration"
    )

    # Save
    output = "Perfect_Foundation_SMS_Report.docx"
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    build_report()
